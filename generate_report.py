from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Inches

def generate_ieee_report():
    doc = Document()
    
    # Setup document defaults (similar to IEEE style)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # TITLE
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Real-Time Multimodal Stress Detection and Analysis using Facial Action Units and Speech Recognition')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # AUTHORS & INSTITUTION
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('Titli Panja, Soumya Ranjan Pradhan\n').font.size = Pt(11)
    inst_run = authors.add_run('IEM CSE(AI)\nSupervisor: Dr. Moutushi Biswas Singh')
    inst_run.italic = True
    inst_run.font.size = Pt(10)
    
    doc.add_paragraph() # spacing

    # ABSTRACT
    abstract_p = doc.add_paragraph()
    abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_title = abstract_p.add_run('Abstract—')
    abs_title.bold = True
    abs_title.font.size = Pt(9)
    abs_body = abstract_p.add_run(
        'Continuous monitoring of psychological stress is critical for mental well-being and mitigating cognitive overload. '
        'In this paper, we present StressAI, a novel multimodal deep learning framework designed to detect human stress '
        'in real-time through simultaneous analysis of facial expressions and speech patterns. Our methodology leverages '
        'geometric extraction of 17 Facial Action Units (AU) and 84 distinct acoustic features (including MFCCs, pitch, '
        'jitter, and shimmer). We propose a hybrid neural architecture comprising a Bidirectional Long Short-Term Memory '
        '(Bi-LSTM) network for visual temporal modeling and a 1D Convolutional Neural Network (CNN) paired with an LSTM '
        'for speech signal processing. Furthermore, a Cross-Modal Attention mechanism is introduced to fuse these streams, '
        'enabling the model to map affective features into the continuous Valence-Arousal (V-A) emotional space. '
        'The model is optimized using a combined Concordance Correlation Coefficient (CCC) and Mean Squared Error (MSE) '
        'loss function. The system features a real-time web-based analytical dashboard deployed via Flask, rendering dynamic '
        'emotional zone classification and stress score timelines.'
    )
    abs_body.font.size = Pt(9)
    abs_body.italic = True

    # INDEX TERMS
    index_p = doc.add_paragraph()
    index_title = index_p.add_run('Index Terms—')
    index_title.bold = True
    index_title.font.size = Pt(9)
    index_body = index_p.add_run('Stress Detection, Affective Computing, Cross-Modal Attention, Facial Action Units, Speech Processing, Valence-Arousal.')
    index_body.font.size = Pt(9)

    doc.add_paragraph()

    # I. INTRODUCTION
    h1 = doc.add_heading('I. INTRODUCTION', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(10)
    h1.runs[0].font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        'In modern environments, chronic psychological stress contributes significantly to adverse cognitive and physical health outcomes. '
        'Traditional clinical evaluations primarily rely on self-reported psychological questionnaires or obtrusive physiological sensors '
        '(such as ECG or GSR), which are subjective or restrictive for everyday continuous monitoring.\n\n'
        'To overcome these limitations, non-contact affective computing has gained traction. However, unimodal systems—relying exclusively '
        'on either visual or acoustic parameters—often struggle with robustness in unconstrained environments due to sensor noise, '
        'lighting variations, and varying microphone qualities.\n\n'
        'This project introduces a highly robust multimodal real-time stress analysis framework. By combining Facial Action Unit (AU) '
        'geometry with advanced acoustic signal processing, the system captures a highly intricate representation of a user’s physiological state. '
        'The primary contribution is a novel end-to-end framework featuring a Cross-Modal Attention deep learning architecture trained on robust '
        'multivariate continuous emotion dimensions (Valence and Arousal) using a hybrid loss metric.'
    )

    # II. PROPOSED METHODOLOGY
    h2 = doc.add_heading('II. PROPOSED METHODOLOGY', level=1)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(10)
    h2.runs[0].font.bold = True

    h2_1 = doc.add_heading('A. Multimodal Feature Extraction Architecture', level=2)
    h2_1.runs[0].font.name = 'Times New Roman'
    h2_1.runs[0].font.size = Pt(10)
    h2_1.runs[0].italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "The system simultaneously captures video and audio streams. The visual pipeline uses the dlib geometric landmark predictor to "
        "estimate 68 distinct facial anchor points. From these topological points, 17 specific anatomical Action Units (such as AU04-Brow Furrow, "
        "and AU12-Lip Corner Pull) are computationally derived based on nodal distances and gradient intensities. Concurrently, the audio pipeline "
        "utilizes the highly robust librosa framework to extract 84 dynamic acoustic metrics, heavily emphasizing micro-prosodic features (jitter, shimmer), "
        "Mel-Frequency Cepstral Coefficients (MFCCs), and harmonic-to-noise ratios (HNR), which are deeply correlated with vocal tract tension under stress."
    )

    h2_2 = doc.add_heading('B. Deep Learning & Cross-Modal Attention', level=2)
    h2_2.runs[0].font.name = 'Times New Roman'
    h2_2.runs[0].font.size = Pt(10)
    h2_2.runs[0].italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "The continuous sequential data is processed through independent temporal networks. The visual domain involves a Bidirectional "
        "LSTM structure engineered to process 30-frame sequence blocks (1 second at 30fps). The acoustic domain uses a serialized 1D CNN "
        "combined with an LSTM to process 20 temporal blocks comprising 500ms intervals.\n\n"
        "Crucially, the independent latent embeddings of these temporal networks are fused using a Cross-Modal Multi-Head Attention layer. "
        "In this layer, the visual latent space establishes attentive weights to the acoustic space and vice versa, mathematically isolating "
        "stress-correlated intersections (e.g., elevated vocal pitch paired inherently with furrowed brows)."
    )

    # III. OPTIMIZATION AND RESULTS
    h3 = doc.add_heading('III. REGRESSION AND OPTIMIZATION', level=1)
    h3.runs[0].font.name = 'Times New Roman'
    h3.runs[0].font.size = Pt(10)
    h3.runs[0].bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Rather than classifying simplistic arbitrary categories, the system utilizes dimensional affect recognition, utilizing the Valence "
        "(pleasantness) and Arousal (physiological activation) circumplex model. The mathematical optimization mitigates standard regression issues "
        "by unifying Mean Squared Error (MSE) and the Concordance Correlation Coefficient (CCC). The final loss function is defined as:\n\n"
        "L = 0.4 × MSE + 0.6 × (1.0 - CCC_avg)\n\n"
        "This dynamic penalty forces the internal neural gradients to aggressively penalize correlation delays, ensuring the real-time "
        "stress timeline heavily matches the user's micro-expression chronologies.\n\n"
        "The following multi-metric analysis visualizes the CCC improvements and correlation scatter across epochs:"
    )

    try:
        from docx.shared import Inches
        doc.add_picture('accuracy_report.png', width=Inches(6.0))
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run('Fig. 1. Validation Concordance Correlation Coefficient vs Epochs')
        img_run.italic = True
        img_run.font.size = Pt(9)
    except Exception as e:
        print(f"Placeholder for image insertion: {e}")

    # IV. SYSTEM ARCHITECTURE & INTERFACE
    h4 = doc.add_heading('IV. SYSTEM ARCHITECTURE', level=1)
    h4.runs[0].font.name = 'Times New Roman'
    h4.runs[0].font.size = Pt(10)
    h4.runs[0].bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "The fully integrated software is deployed locally as an asynchronous Flask web server with a heavily optimized multi-threaded "
        "inference pipeline. The graphical interface translates V-A arrays into highly interpretable analytical reports, visualizing metrics "
        "such as high-stress percentages, baseline deviations, and quadrant-based emotional zoning (e.g., Anxiety versus Anger). It also provides "
        "context-aware actionable recommendations driven by the specific trajectory of the session."
    )

    # V. CONCLUSION
    h5 = doc.add_heading('V. CONCLUSION', level=1)
    h5.runs[0].font.name = 'Times New Roman'
    h5.runs[0].font.size = Pt(10)
    h5.runs[0].bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "The StressAI model presents an intricate and highly responsive system for real-time psychoanalytical monitoring. "
        "By synthesizing geometric visual features with acoustic variance mechanisms via advanced Cross-Modal Neural Attention, "
        "the application mitigates environmental noise constraints and produces highly reliable Valence-Arousal indices. "
        "Future iterations will involve expanding the physiological architecture to include remote Photoplethysmography (rPPG) "
        "for purely visual heart-rate derivation."
    )

    # REFERENCES
    h6 = doc.add_heading('REFERENCES', level=1)
    h6.runs[0].font.name = 'Times New Roman'
    h6.runs[0].font.size = Pt(10)
    h6.runs[0].bold = True
    
    ref = doc.add_paragraph()
    ref.style = doc.styles['Normal']
    ref.add_run("[1] P. Ekman and W. V. Friesen, “Facial Action Coding System: A Technique for the Measurement of Facial Movement,” Consulting Psychologists Press, 1978.\n").font.size = Pt(9)
    ref.add_run("[2] J. A. Russell, “A circumplex model of affect,” Journal of Personality and Social Psychology, vol. 39, pp. 1161-1178, 1980.\n").font.size = Pt(9)
    ref.add_run("[3] Ashish Vaswani et al., \"Attention is all you need,\" in Advances in Neural Information Processing Systems, 2017.\n").font.size = Pt(9)
    ref.add_run("[4] Lawrence Lin, “A Concordance Correlation Coefficient to Evaluate Reproducibility,” Biometrics, vol. 45, no. 1, 1989.\n").font.size = Pt(9)

    doc.save('Stress_Analysis_IEEE_Report.docx')
    print("Document successfully generated at Stress_Analysis_IEEE_Report.docx")

if __name__ == '__main__':
    generate_ieee_report()
